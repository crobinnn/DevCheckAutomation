import websockets
import asyncio
import json
import base64
import openpyxl
import numpy as np
import cv2
import os
import openpyxl
import threading as Thread
from openpyxl.drawing.image import Image as ExcelImage
from io import BytesIO
from PIL import Image as PILImage
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from pyzbar.pyzbar import decode as zdecode, ZBarSymbol

# Define a class for managing shared data for each connection
class ConnectionData:
    def __init__(self):
        self.serial_length = 10
        self.mappings = []
  
def create_documents_from_mappings(excel_template_path, excel_output_filename, word_output_filename, mappings):
  cwd = os.getcwd()
  
  excel_output_path = os.path.join(cwd, excel_output_filename)
  word_output_path = os.path.join(cwd, word_output_filename)
  
  wb = openpyxl.load_workbook(excel_template_path)
  sheet_name = 'Sheet1'
  if sheet_name in wb.sheetnames:
    ws = wb[sheet_name]
  else:
    ws = wb.active  

  # Cell Excel E12 SN, F12 MAC, I12 img
  start_row = 12
  serial_col = 5 # Column E
  mac_col = 6  # Column F

  for i, mapping in enumerate(mappings):
    serial_numbers = mapping["serial_number"]
    mac_addresses = mapping["mac_address"]
    
    serial_cell = ws.cell(row=start_row + i, column=serial_col)
    mac_cell = ws.cell(row=start_row + i, column=mac_col)

    serial_cell.value = ", ".join(serial_numbers) if serial_numbers else ""
    mac_cell.value = ", ".join(mac_addresses) if mac_addresses else ""

  wb.save(excel_output_path)
  print(f'created excel at {excel_output_path}')
  
  document = Document()
  document.add_heading('Images and Serial Numbers', level=1)
  
  table = document.add_table(rows=0, cols=2)
  table.autofit = True
  table.style = 'Table Grid'
  
  for i in range(0, len(mappings), 2): 
    row_cells = table.add_row().cells

    # Add a new row for each image and its serial number
    for col, mapping in enumerate(mappings[i:i+2]):
      # Decode base64 to image
      img_data = base64.b64decode(mapping["image_data"])
      img_pil = PILImage.fromarray(cv2.cvtColor(cv2.imdecode(np.frombuffer(img_data, np.uint8), cv2.IMREAD_COLOR), cv2.COLOR_BGR2RGB))
      img_byte_arr = BytesIO()
      img_pil.save(img_byte_arr, format='PNG')
      img_byte_arr.seek(0)

      # Add img to cell
      paragraph = row_cells[col].paragraphs[0]
      run = paragraph.add_run()
      run.add_picture(img_byte_arr, width=Inches(2.0))

      # Add SN below img
      row_cells[col].add_paragraph(f'SN: {", ".join(mapping["serial_number"])}')

  # Save Word
  document.save(word_output_path)
  print(f'created word at {word_output_path}')

# Image Processing
def process_image(image_bytes,connection_data):
    image_np = np.frombuffer(image_bytes, np.uint8)
    image = cv2.imdecode(image_np, cv2.IMREAD_COLOR)

    # Change image to grayscale
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    max_value = gray.max()
    print(f"Max grayscale value: {max_value}")

    # Calculate contrast factor
    contrast_factor = 255 / max_value if max_value > 0 else 1
    contrasted = cv2.multiply(gray, contrast_factor)

    # Apply thresholding
    ret, thresh = cv2.threshold(contrasted, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

    # Invert colors for barcode recognition
    inverted = cv2.bitwise_not(thresh)
    restored = cv2.bitwise_not(inverted)

    decoded_objects = zdecode(restored, symbols=[ZBarSymbol.QRCODE, ZBarSymbol.CODE128])

    serial_numbers = []
    mac_addresses = []

    for obj in decoded_objects:
        data = obj.data.decode('utf-8')
        print(f"Decoded data: {data}")  

        expected_length = int(connection_data.serial_length)

        if len(data) == expected_length or (len(data) == 19 and '-' in data):
            serial_numbers.append(data)
        elif ((len(data) == 14 and all(c in '0123456789ABCDEF:' for c in data) and ':' in data) or
          ((len(data) == 17 and all(c in '0123456789ABCDEF:' for c in data) and ':' in data) or
          (len(data) == 12 and all(c in '0123456789ABCDEF' for c in data)))):
            mac_addresses.append(data)

    return serial_numbers, mac_addresses

async def serial_length_node(websocket, message_str,connection_data):
    try:
        data = json.loads(message_str) 
        if "serial_length" in data:
            connection_data.serial_length = data["serial_length"]
            print(f"Updated serial length to: {connection_data.serial_length}")
            await websocket.send(json.dumps({"status": "Serial length updated"}))
    except Exception as e:
        print(f"Error processing serial length message: {e}")

async def image_processing_node(websocket, image_data,connection_data):
    try:
        serial_numbers, mac_addresses = process_image(image_data,connection_data)

        # Encode image to base64
        image_base64 = base64.b64encode(image_data).decode('utf-8')

        # Store the mapping in the shared data
        connection_data.mappings.append({
            "image_data": image_base64,  # Store Base64 string
            "serial_number": serial_numbers,
            "mac_address": mac_addresses
        })

        # Print each processed serial number and MAC address
        if serial_numbers:
            print(f"Processed Serial Numbers: {', '.join(serial_numbers)}")
        else:
            print("No serial numbers detected.")

        if mac_addresses:
            print(f"Processed MAC Addresses: {', '.join(mac_addresses)}")
        else:
            print("No MAC addresses detected.")

        response_data = {
            "serial_number": serial_numbers,
            "mac_address": mac_addresses,
            "image_data": image_base64  # Send the Base64 encoded image
        }

        await websocket.send(json.dumps(response_data))
        print(f"Sent response: {{'serial_number': {serial_numbers}, 'mac_address': {mac_addresses}, 'image_data': 'image received'}}")
    except Exception as e:
        print(f"Error processing image data: {e}")
        
async def get_mappings_node(websocket,connection_data):
    try:
        # Send the current mappings back to the client with Base64 images
        await websocket.send(json.dumps({"mappings": connection_data.mappings}))
        print(f"Sent mappings data to client: {connection_data.mappings}")
    except Exception as e:
        print(f"Error sending mappings data: {e}")

async def delete_mapping_node(websocket, identifier,connection_data):
    try:
        # Remove mapping based on identifier (either index or serial number)
        if isinstance(identifier, int) and 0 <= identifier < len(connection_data.mappings):
            deleted_mapping = connection_data.mappings.pop(identifier)
            print(f"Deleted mapping at index {identifier}: {deleted_mapping}")
            await websocket.send(json.dumps({"status": "Mapping deleted", "index": identifier}))
        else:
            await websocket.send(json.dumps({"status": "Invalid identifier"}))
    except Exception as e:
        print(f"Error deleting mapping: {e}")
        await websocket.send(json.dumps({"status": "Error deleting mapping"}))

async def create_document_node(websocket, data,connection_data):
    print(f"Received data for document creation: {data}")
    try:
        if "create_documents" in data:
            nested_data = data["create_documents"]
            excel_template_path = 'C:\\automationcompnet\\automat\\EXCEL MAKING\\1Template.xlsx'
            excel_output_filename = nested_data.get("excel_file")
            word_output_filename = nested_data.get("word_file")
        if not excel_output_filename or not word_output_filename:
          print("Error: Missing filename(s) for output.")
          await websocket.send(json.dumps({"status": "Error: Missing filename(s) for output."}))
          return
        create_documents_from_mappings(excel_template_path, excel_output_filename, word_output_filename, connection_data.mappings)
        print("Documents created successfully")
        await websocket.send(json.dumps({
            "status": "Documents created",
            "excel_file": excel_output_filename,
            "word_file": word_output_filename
        }))
    except Exception as e:
        print(f"Error creating documents: {e}")
        await websocket.send(json.dumps({"status": "Error creating documents"}))
        
async def server_handler(websocket):
    connection_data = ConnectionData()  # Each connection has its own shared data
    client_ip, client_port = websocket.remote_address
    print(f"Client connected from IP: {client_ip}, Port: {client_port}")
    try:
        async for message in websocket:
            # Check if the message is a byte stream (for images)
            if isinstance(message, bytes):  
                await image_processing_node(websocket, message, connection_data)
            else:
                # If it's not bytes, handle it as a JSON string directly
                try:
                    data = json.loads(message)
                    if "get_mappings" in data:
                        await get_mappings_node(websocket, connection_data)
                    elif "delete_mapping" in data:
                        identifier = data["delete_mapping"]
                        await delete_mapping_node(websocket, identifier, connection_data)
                    elif "create_documents" in data:
                        await create_document_node(websocket, data, connection_data)
                    else:
                        await serial_length_node(websocket, message, connection_data)
                except Exception as e:
                    print(f"Error processing message: {e}")
    except Exception as e:
        print(f"Error: {e}")
    finally:
        await websocket.close()

# Setup WebSocket server
async def main():
    async with websockets.serve(server_handler, "0.0.0.0", 8765):
        await asyncio.Future()  # run forever

if __name__ == "__main__":
    asyncio.run(main())


excel_path = 'C:\\automationcompnet\\automat\\EXCEL MAKING\\1Template.xlsx'