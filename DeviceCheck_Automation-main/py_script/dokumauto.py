import cv2
import numpy as np
from pyzbar.pyzbar import decode as zdecode, ZBarSymbol
import os
import openpyxl
from openpyxl.drawing.image import Image as ExcelImage
from io import BytesIO
from PIL import Image as PILImage
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

def process_image(image_path):
  # Load the image
  image = cv2.imread(image_path)
  # Change img grayscale
  gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
  max_value = gray.max()
  print(max_value)
  
  # Calculate contrast factor
  contrast_factor = 255 / max_value if max_value > 0 else 1
  
  contrasted = cv2.multiply(gray, contrast_factor)
  
  # Apply thresholding
  ret, thresh = cv2.threshold(contrasted, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
  # Invert color barcode jd putih bg hitam
  inverted = cv2.bitwise_not(thresh)
  # Invert lagi biar jd item lg barcode biar lebi clear
  restored = cv2.bitwise_not(inverted)

  # # Gaussian Blur minimalisir noise
  # blur = cv2.GaussianBlur(gray, (5, 5), 0)
  # # Contrast Limited Adaptive Histogram Equalization
  # clahe = cv2.createCLAHE(clipLimit=5.0, tileGridSize=(48, 48))
  # contrast = clahe.apply(blur)
  # # OTSU's Method Tresholding
  # ret, thresh = cv2.threshold(contrast, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
  # # Invert color barcode jd putih bg hitam
  # inverted = cv2.bitwise_not(thresh)
  # # Invert lagi biar jd item lg barcode biar lebi clear
  # restored = cv2.bitwise_not(inverted)

  decoded_objects = zdecode(restored, symbols=[ZBarSymbol.QRCODE, ZBarSymbol.CODE128])
  
  serial_number = None
  mac_address = None
  serial_numbers = []
  mac_addressu = []

  # Get Data barcode
  for obj in decoded_objects:
    data = obj.data.decode('utf-8')
    if (len(data) == 12 or 
      (len(data) == 11 ) or
      (len(data) == 16) or
      # # (len(data) == 27) or
      (len(data) == 19 and '-' in data)):
      serial_numbers.append(data)
    elif ((len(data) == 14 and all(c in '0123456789ABCDEF:' for c in data) and ':' in data) or
          ((len(data) == 17 and all(c in '0123456789ABCDEF:' for c in data) and ':' in data) or
      (len(data) == 12 and all(c in '0123456789ABCDEF' for c in data)))):
      mac_addressu.append(data)

  return serial_numbers, mac_addressu

def resize_image(image_path, size):
  image = cv2.imread(image_path)
  def unsharp_mask(resized_image, kernel_size=(5, 5), sigma=0.5, amount=1.0, threshold=0):
    """Return a sharpened version of the image, using an unsharp mask."""
    blurred = cv2.GaussianBlur(resized_image, kernel_size, sigma)
    sharpened = float(amount + 1) * resized_image - float(amount) * blurred
    sharpened = np.maximum(sharpened, np.zeros(sharpened.shape))
    sharpened = np.minimum(sharpened, 255 * np.ones(sharpened.shape))
    sharpened = sharpened.round().astype(np.uint8)
    if threshold > 0:
        low_contrast_mask = np.absolute(resized_image - blurred) < threshold
        np.copyto(sharpened, resized_image, where=low_contrast_mask)
    return sharpened

  sharpening_kernel = np.array([
    [0, -0.4, 0],
    [-0.4, 2.6, -0.4],
    [0, -0.4, 0]
  ])

  sharpened_image = cv2.filter2D(image, -1, sharpening_kernel)

  sharpened_image_rgb = cv2.cvtColor(sharpened_image, cv2.COLOR_BGR2RGB)

  pil_image = PILImage.fromarray(sharpened_image_rgb)

  # Resize image
  resized_pil_image = pil_image.resize(size, PILImage.LANCZOS)

  resized_image = cv2.cvtColor(np.array(resized_pil_image), cv2.COLOR_RGB2BGR)

  final_image = unsharp_mask(resized_image)
  # resized_image = cv2.resize(image, size, interpolation=cv2.INTER_LANCZOS4)
  return final_image
  
def write_to_excel(data, excel_path, sheet_name, images):
  wb = openpyxl.load_workbook(excel_path)
  if sheet_name in wb.sheetnames:
    ws = wb[sheet_name]
  else:
    ws = wb.active  

  # Cell Excel E12 SN, F12 MAC, I12 img
  start_row = 12
  serial_col = 5 # Column E
  mac_col = 6  # Column F
  img_col = 9 # Column I

  for i, (serial, mac) in enumerate(data):
    serial_cell = ws.cell(row=start_row + i, column=serial_col)
    mac_cell = ws.cell(row=start_row + i, column=mac_col)
    if serial and mac:
      serial_cell.value = ", ".join(serial)
      mac_cell.value = ", ".join(mac)
    elif serial:
      serial_cell.value = ", ".join(serial)
      mac_cell.value = ""
    elif mac:
      serial_cell.value = ""
      mac_cell.value = ", ".join(mac)
    else:
      serial_cell.value = ""
      mac_cell.value = ""
    
    # Convert image to in-memory file for Excel insertion
    img = PILImage.fromarray(cv2.cvtColor(images[i], cv2.COLOR_BGR2RGB))
    img_byte_arr = BytesIO()
    img.save(img_byte_arr, format='PNG')
    img_byte_arr.seek(0)
    img_excel = ExcelImage(img_byte_arr)
    
    img_excel.anchor = ws.cell(row=start_row + i, column=img_col).coordinate
    ws.add_image(img_excel)
    pixel_width, pixel_height = img.size
    points_width = pixel_width * 0.75
    points_height = pixel_height * 0.75
    ws.row_dimensions[start_row + i].height = points_height
    ws.column_dimensions[openpyxl.utils.get_column_letter(img_col)].width = points_width / 7  # Approximation

  wb.save(excel_path)

def create_word_document(excel_path, images, serial_numbers, word_output_path):
    document = Document()
    document.add_heading('Images and Serial Numbers', level=1)

    # Create a table with two columns
    table = document.add_table(rows=0, cols=2)
    table.autofit = True
    table.style = 'Table Grid'

    # Iterate over the images and serial numbers together
    for i in range(0, len(images), 2):
        # Add a new row for each pair of images
        row_cells = table.add_row().cells

        # First column
        if i < len(images):
            img_pil = PILImage.fromarray(cv2.cvtColor(images[i], cv2.COLOR_BGR2RGB))
            img_byte_arr = BytesIO()
            img_pil.save(img_byte_arr, format='PNG')
            img_byte_arr.seek(0)
            row_cells[0].paragraphs[0].add_run().add_picture(img_byte_arr, width=Inches(2.0))
            row_cells[0].add_paragraph(f'SN: {", ".join(serial_numbers[i])}')

        # Second column (if there's a second image)
        if i + 1 < len(images):
            img_pil = PILImage.fromarray(cv2.cvtColor(images[i + 1], cv2.COLOR_BGR2RGB))
            img_byte_arr = BytesIO()
            img_pil.save(img_byte_arr, format='PNG')
            img_byte_arr.seek(0)
            row_cells[1].paragraphs[0].add_run().add_picture(img_byte_arr, width=Inches(2.0))
            row_cells[1].add_paragraph(f'SN: {", ".join(serial_numbers[i + 1])}')

    document.save(word_output_path)


def process_folder(folder_path, excel_path, sheet_name):
  data = []
  images = []
  image_files = [f for f in os.listdir(folder_path) if f.lower().endswith(('.png', '.jpg', '.jpeg'))]
    
  num_images = len(image_files)
  print(f"Number of images found: {num_images}")
  for filename in os.listdir(folder_path):
    if filename.lower().endswith(('.png', '.jpg', '.jpeg')):
      image_path = os.path.join(folder_path, filename)
      serial, mac = process_image(image_path)
      data.append((serial, mac))
      if serial and mac:
        print(f"Image: {filename} - Serial Number: {serial}, MAC Address: {mac}")
      elif serial:
        print(f"Image: {filename} - Serial Number: {serial} added, but MAC Address is not readable.")
      elif mac:
        print(f"Image: {filename} - MAC Address: {mac} added, but Serial Number is not readable.")
      else:
        print(f"Image: {filename} - No barcode detected")
    
      # # Resize timage
      final_image = resize_image(image_path, size=(400, 400))
      images.append(final_image)

  write_to_excel(data, excel_path, sheet_name, images)
  print("Process Completed", "Barcode data has been added to the Excel file.")
  create_word_document(excel_path, images, [serial for serial, _ in data], word_output_path)

# Folder gambar
folder_path = 'C:\\automationcompnet\\automat\\EXCEL MAKING\\foto'

# Output Excel file path
excel_path = 'C:\\automationcompnet\\automat\\EXCEL MAKING\\Forti_Wisnu.xlsx'

word_output_path = 'C:\\automationcompnet\\automat\\EXCEL MAKING\\Forti_Wisnu.docx'
# Sheet excel
sheet_name = 'Sheet1'

process_folder(folder_path, excel_path, sheet_name)